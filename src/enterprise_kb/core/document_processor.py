"""Document processor module — parses PDF, Word, and Markdown documents.

Supports:
- LlamaParse (cloud) for complex PDF parsing
- python-docx for Word documents (fallback)
- Built-in Markdown parser
- Batch processing with async IO
"""

import io
import mimetypes
from pathlib import Path
from typing import Any, AsyncIterator, Optional

import aiofiles
import httpx
import markdown
from bs4 import BeautifulSoup

from enterprise_kb.config import settings
from enterprise_kb.utils.logger import logger


class DocumentProcessingError(Exception):
    """Raised when document parsing fails."""


class Document:
    """Represents a parsed document with metadata."""

    def __init__(
        self,
        content: str,
        metadata: dict[str, Any],
        file_path: str,
        file_type: str,
    ) -> None:
        self.content = content
        self.metadata = metadata
        self.file_path = file_path
        self.file_type = file_type

    def __repr__(self) -> str:
        return f"Document(path={self.file_path}, type={self.file_type}, len={len(self.content)})"


class DocumentProcessor:
    """Processes and parses documents from various file formats.

    Parsing strategy (in order of preference):
    1. LlamaParse API (if configured) for PDF
    2. python-docx for Word (.docx)
    3. Built-in Markdown parser
    """

    def __init__(self) -> None:
        self._http_client: httpx.AsyncClient | None = None
        self._llama_available = settings.use_llama_parse
        self._llama_api_key = settings.llama_cloud_api_key

    async def _ensure_client(self) -> httpx.AsyncClient:
        if self._http_client is None:
            self._http_client = httpx.AsyncClient(timeout=120.0)
        return self._http_client

    async def process_file(self, file_path: str | Path) -> Document:
        """Parse a single file based on its extension.

        Args:
            file_path: Path to the document (PDF, .docx, .md).

        Returns:
            A :class:`Document` with parsed content and metadata.

        Raises:
            DocumentProcessingError: If parsing fails.
            FileNotFoundError: If the file does not exist.
        """
        path = Path(file_path)
        if not path.is_file():
            raise FileNotFoundError(f"File not found: {path}")

        ext = path.suffix.lower()
        stats = path.stat()

        base_meta: dict[str, Any] = {
            "file_name": path.name,
            "file_path": str(path.resolve()),
            "file_size": stats.st_size,
            "file_type": ext,
            "last_modified": stats.st_mtime,
        }

        if ext == ".pdf":
            content = await self._parse_pdf(path)
        elif ext == ".docx":
            content = await self._parse_docx(path)
        elif ext in (".md", ".markdown"):
            content = await self._parse_markdown(path)
        else:
            # Try to detect mime type
            mime_type, _ = mimetypes.guess_type(str(path))
            if mime_type and "text" in mime_type:
                content = await self._parse_plain_text(path)
            else:
                msg = f"Unsupported file type: {ext}"
                raise DocumentProcessingError(msg)

        return Document(
            content=content,
            metadata=base_meta,
            file_path=str(path),
            file_type=ext,
        )

    async def process_batch(
        self,
        file_paths: list[str | Path],
        concurrency: int = 4,
    ) -> list[Document]:
        """Parse multiple files concurrently.

        Args:
            file_paths: List of file paths to process.
            concurrency: Maximum number of concurrent parses.

        Returns:
            List of parsed :class:`Document` objects.

        Raises:
            DocumentProcessingError: If all files fail.
        """
        import asyncio

        sem = asyncio.Semaphore(concurrency)

        async def _limited(path: str | Path) -> Document | None:
            async with sem:
                try:
                    return await self.process_file(path)
                except Exception:
                    logger.exception("Failed to parse %s", path)
                    return None

        tasks = [_limited(p) for p in file_paths]
        results = await asyncio.gather(*tasks)
        docs = [r for r in results if r is not None]

        if not docs:
            msg = "All files failed to parse"
            raise DocumentProcessingError(msg)

        return docs

    async def process_directory(
        self,
        directory: str | Path,
        pattern: str = "*.md",
        recursive: bool = True,
    ) -> list[Document]:
        """Parse all files matching a pattern in a directory.

        Args:
            directory: Root directory to scan.
            pattern: Glob pattern (default: ``*.md``).
            recursive: Search recursively.

        Returns:
            List of parsed :class:`Document` objects.
        """
        root = Path(directory)
        if not root.is_dir():
            raise NotADirectoryError(f"Not a directory: {root}")

        glob = root.rglob if recursive else root.glob
        paths = [p for p in glob(pattern) if p.is_file()]
        return await self.process_batch(paths)

    # ── Format-specific parsers ──

    async def _parse_pdf(self, path: Path) -> str:
        """Parse PDF using LlamaParse (cloud) or fallback."""
        if self._llama_available:
            return await self._parse_pdf_llama(path)

        # Fallback: try PyMuPDF / pdfminer if available
        try:
            return await self._parse_pdf_fallback(path)
        except Exception:
            logger.warning("No PDF parser available. Install `pypdf` or configure LlamaParse.")
            msg = "No PDF parser configured"
            raise DocumentProcessingError(msg) from None

    async def _parse_pdf_llama(self, path: Path) -> str:
        """Parse PDF via LlamaParse API."""
        client = await self._ensure_client()
        files = {"file": (path.name, path.read_bytes(), "application/pdf")}

        headers = {"Authorization": f"Bearer {self._llama_api_key}"}
        resp = await client.post(
            "https://api.cloud.llamaindex.ai/api/v1/parsing/upload",
            files=files,
            headers=headers,
        )
        resp.raise_for_status()
        data = resp.json()

        # LlamaParse returns job_id — poll for result
        job_id = data.get("id", data.get("job_id"))
        if not job_id:
            raise DocumentProcessingError("LlamaParse did not return a job ID")

        result = await self._poll_llama_job(client, job_id, headers)
        return result

    async def _poll_llama_job(
        self,
        client: httpx.AsyncClient,
        job_id: str,
        headers: dict[str, str],
        max_retries: int = 30,
        interval: float = 2.0,
    ) -> str:
        """Poll LlamaParse job until completion."""
        import asyncio

        url = f"https://api.cloud.llamaindex.ai/api/v1/parsing/job/{job_id}"

        for attempt in range(max_retries):
            resp = await client.get(url, headers=headers)
            resp.raise_for_status()
            data = resp.json()

            status = data.get("status", "").lower()
            if status == "completed":
                result_url = data.get("result_url") or data.get("url", "")
                if result_url:
                    text_resp = await client.get(result_url)
                    text_resp.raise_for_status()
                    return text_resp.text
                # Result inline
                return data.get("result", {}).get("text", "")
            elif status in ("error", "failed"):
                err_msg = data.get("error", "Unknown LlamaParse error")
                raise DocumentProcessingError(f"LlamaParse failed: {err_msg}")

            await asyncio.sleep(interval)

        raise DocumentProcessingError("LlamaParse timed out")

    async def _parse_pdf_fallback(self, path: Path) -> str:
        """Fallback PDF parser using pypdf."""
        import pypdf

        reader = pypdf.PdfReader(str(path))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n\n".join(pages)

    async def _parse_docx(self, path: Path) -> str:
        """Parse .docx file using python-docx."""
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text:
                    paragraphs.append(row_text)

        return "\n\n".join(paragraphs)

    async def _parse_markdown(self, path: Path) -> str:
        """Parse .md file, strip HTML, return plain text."""
        raw = await self._read_file(path)
        html = markdown.markdown(raw, extensions=["fenced_code", "tables"])
        soup = BeautifulSoup(html, "html.parser")
        return soup.get_text(separator="\n", strip=True)

    async def _parse_plain_text(self, path: Path) -> str:
        """Read plain text file."""
        return await self._read_file(path)

    @staticmethod
    async def _read_file(path: Path) -> str:
        """Read a UTF-8 text file."""
        async with aiofiles.open(path, encoding="utf-8") as f:
            return await f.read()

    async def close(self) -> None:
        """Release HTTP client resources."""
        if self._http_client:
            await self._http_client.aclose()
            self._http_client = None
